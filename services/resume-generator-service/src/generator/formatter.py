from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import logging
from typing import Dict, List
import tempfile

logger = logging.getLogger(__name__)

class ResumeFormatter:
    def __init__(self):
        """Initialize the resume formatter"""
        self.default_template_path = os.path.join(
            os.path.dirname(__file__),
            "template.docx"
        )

    def create_document(
        self,
        content: Dict,
        template_path: str = None
    ) -> str:
        """
        Create a formatted resume document
        Args:
            content: Dictionary containing resume content
            template_path: Optional path to custom template
        Returns:
            Path to the generated document
        """
        try:
            # Use template if provided and exists, otherwise create new
            if template_path and os.path.exists(template_path):
                return self._apply_template(content, template_path)
            else:
                return self._create_new_document(content)
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            raise

    def _apply_template(self, content: Dict, template_path: str) -> str:
        """Apply content to existing template"""
        try:
            doc = DocxTemplate(template_path)
            
            # Prepare context for template
            context = {
                "summary": content.get("summary", ""),
                "experience": content.get("experience", []),
                "skills": ", ".join(content.get("skills", [])),
                "metadata": content.get("metadata", {})
            }
            
            # Render template
            doc.render(context)
            
            # Save to temporary file
            output_path = tempfile.mktemp(suffix=".docx")
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error applying template: {e}")
            raise

    def _create_new_document(self, content: Dict) -> str:
        """Create new document from scratch"""
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Add metadata
            metadata = content.get("metadata", {})
            if metadata:
                heading = doc.add_heading(level=1)
                name_run = heading.add_run("Resume")
                name_run.font.size = Pt(18)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add job details
                if metadata.get("job_title") and metadata.get("company"):
                    subheading = doc.add_paragraph()
                    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    subheading.add_run(
                        f"For {metadata['job_title']} at {metadata['company']}"
                    ).font.size = Pt(12)
            
            # Add summary
            if content.get("summary"):
                doc.add_heading("Professional Summary", level=2)
                summary_para = doc.add_paragraph()
                summary_para.add_run(content["summary"]).font.size = Pt(11)
            
            # Add experience
            if content.get("experience"):
                doc.add_heading("Experience", level=2)
                for entry in content["experience"]:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(entry["description"]).font.size = Pt(11)
            
            # Add skills
            if content.get("skills"):
                doc.add_heading("Skills", level=2)
                skills_para = doc.add_paragraph()
                skills_para.add_run(", ".join(content["skills"])).font.size = Pt(11)
            
            # Save to temporary file
            output_path = tempfile.mktemp(suffix=".docx")
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating new document: {e}")
            raise

    def _format_paragraph(self, paragraph, font_size: int = 11):
        """Apply consistent formatting to a paragraph"""
        paragraph.style = 'Normal'
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.name = 'Calibri'
